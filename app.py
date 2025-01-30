from flask import Flask, request, render_template, send_file, redirect, url_for
from google.cloud import firestore
import os
import json
import base64
from docx import Document
from datetime import datetime, timedelta
import firebase_admin
from firebase_admin import credentials, firestore

app = Flask(__name__)


# # Set the Firebase service account key
# os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "credentials.json"

# # Initialize Firestore client
# db = firestore.Client()

# Load Firebase Credentials from Environment Variables
firebase_key_base64 = os.getenv("FIREBASE_KEY")
if not firebase_key_base64:
    raise ValueError("FIREBASE_KEY environment variable is not set")

firebase_key = json.loads(base64.b64decode(firebase_key_base64).decode("utf-8"))
cred = credentials.Certificate(firebase_key)
firebase_admin.initialize_app(cred)

# Initialize Firestore
db = firestore.client()

# Helper function to replace text in Word document
def replace_text_in_run(run, key, value):
    if key in run.text:
        run.text = run.text.replace(key, value)
        run.font.bold = True

def replace_placeholders(doc, placeholders):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in placeholders.items():
                replace_text_in_run(run, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in placeholders.items():
                            replace_text_in_run(run, key, value)

def create_bond_letter(data, output_folder, unique_id):
    # Load the bond template
    template_path = "BOND.DOCX"
    doc = Document(template_path)

    # Compute Joining Date (exactly 1 year after "Date")
    date_selected = datetime.strptime(data.get("date"), "%Y-%m-%d")
    end_date = date_selected + timedelta(days=365)

    # Prepare placeholders
    placeholders = {
        "[NAME]": data.get("name", "N/A"),
        "MOBILE": data.get("mobile", "N/A"),
        "DATE": date_selected.strftime("%d/%m/%Y"),
        "[ADDRESS]": data.get("address", "N/A"),
        "AADHAR": data.get("aadhar", "N/A"),
        "<ROLE>": data.get("role", "N/A"),
        "JODA": end_date.strftime("%d/%m/%Y"),  # Auto-calculated joining date
    }

    # Replace placeholders in the document
    replace_placeholders(doc, placeholders)

    # Create output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Save the bond letter
    file_name = f"BOND_HETERIZE_INFOTECH_{data.get('name', 'N/A')}_{data.get('mobile', 'N/A')}.docx"
    file_path = os.path.join(output_folder, file_name)
    doc.save(file_path)

    return file_path

def save_to_firestore(data):
    """Save the bond details into Firestore."""
    collection_name = "BOND"

    # Compute Joining Date (exactly 1 year after "Date")
    date_selected = datetime.strptime(data.get("date"), "%Y-%m-%d")
    end_date = date_selected + timedelta(days=365)
    data["end_date"] = end_date.strftime("%Y-%m-%d")  # Save in Firestore

    # Add a timestamp
    data["timestamp"] = datetime.utcnow().isoformat()

    # Save the data in Firestore with a unique document ID
    doc_ref = db.collection(collection_name).document()
    unique_id = doc_ref.id  # Firestore-generated unique ID
    data["unique_id"] = unique_id  # Add unique ID to the data

    doc_ref.set(data)
    return unique_id

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Extract form data
        data = request.form.to_dict()

        # Ensure required fields are present
        required_fields = ["name", "mobile", "date", "address", "aadhar", "role"]
        for field in required_fields:
            if field not in data or not data[field]:
                return f"Missing required field: {field}", 400

        # Save the data to Firestore and get the unique ID
        try:
            unique_id = save_to_firestore(data)

            # Generate the bond letter
            output_folder = "bond_letters"
            file_path = create_bond_letter(data, output_folder, unique_id)

            # Redirect to download the file
            return redirect(url_for("download", file_name=os.path.basename(file_path)))

        except Exception as e:
            return f"An error occurred: {e}", 500

    return render_template("index.html")

@app.route("/download/<file_name>")
def download(file_name):
    file_path = os.path.join("bond_letters", file_name)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return "File not found", 404

if __name__ == "__main__":
    app.run(debug=True)
